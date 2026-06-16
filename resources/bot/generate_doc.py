"""Generate comprehensive Word document for the WeChat bot project."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, code_text):
    """Add a styled code block paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Code']
    p.text = code_text
    return p


def add_heading_with_line(doc, text, level=1):
    """Add heading with a separator line."""
    heading = doc.add_heading(text, level=level)
    return heading


def create_styles(doc):
    """Create custom styles for the document."""
    # Code style
    style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Consolas'
    style.font.size = Pt(8.5)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.left_indent = Cm(0.5)

    # Inline code character style
    style = doc.styles.add_style('InlineCode', WD_STYLE_TYPE.CHARACTER)
    style.font.name = 'Consolas'
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

    # Note box style
    style = doc.styles.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(10)
    style.font.italic = True
    style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    style.paragraph_format.left_indent = Cm(1)


def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table


def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    create_styles(doc)

    # ══════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("赵有才微信机器人")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("WeChat Auto-Reply Bot\n完整开发文档与使用说明书")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.date.today().strftime("%Y年%m月%d日")
    run = info.add_run(f"版本 2.0  |  {today}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info2.add_run("基于 Claude API + WeChat 4.x + WeFlow Electron 集成")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("目录", level=1)

    toc_items = [
        ("1.", "项目概述", "3"),
        ("2.", "系统架构总览", "4"),
        ("2.1", "整体数据流", "4"),
        ("2.2", "三大运行模式", "5"),
        ("2.3", "技术栈总览", "6"),
        ("3.", "Python 后端模块详解", "7"),
        ("3.1", "入口调度 — main.py", "7"),
        ("3.2", "核心引擎 — bot_image.py", "8"),
        ("3.3", "轻量引擎 — bot_reply_server.py", "10"),
        ("3.4", "WeChatFerry 后端 — bot_wcf.py", "11"),
        ("3.5", "统一发送模块 — send_backend.py", "12"),
        ("3.6", "窗口操作模块 — wechat_send.py", "13"),
        ("3.7", "人类化输入 — human_send.py", "14"),
        ("3.8", "AI 客户端 — claude_client.py", "15"),
        ("3.9", "对话管理 — conversation_manager.py", "16"),
        ("3.10", "联系人图谱 — contacts.py", "16"),
        ("3.11", "人设提示词 — zhaoyoucai_prompt.py", "17"),
        ("3.12", "OCR 引擎 — ocr_engine.py", "18"),
        ("3.13", "WebSocket 服务 — ws_server.py", "19"),
        ("3.14", "终端显示 — display.py", "19"),
        ("3.15", "配置中心 — config.py", "20"),
        ("3.16", "辅助工具模块", "20"),
        ("4.", "Electron 前端详解", "21"),
        ("4.1", "主进程 — main.ts", "21"),
        ("4.2", "桥接服务 — botBridgeService.ts", "22"),
        ("4.3", "预加载脚本 — preload.ts", "23"),
        ("4.4", "Bot 管理页 — BotPage.tsx", "23"),
        ("4.5", "回复输入栏 — BotReplyBar.tsx", "24"),
        ("4.6", "状态管理 — botStore.ts", "24"),
        ("5.", "AI 智能回复系统", "25"),
        ("5.1", "多维度上下文构建", "25"),
        ("5.2", "人设系统设计", "26"),
        ("5.3", "对话历史管理", "27"),
        ("5.4", "API 调用策略", "27"),
        ("6.", "消息发送机制", "28"),
        ("6.1", "ClipboardBackend — 剪贴板模式", "28"),
        ("6.2", "WcfBackend — WeChatFerry 模式", "29"),
        ("6.3", "发送流程对比", "29"),
        ("7.", "用户使用手册", "30"),
        ("7.1", "环境要求", "30"),
        ("7.2", "安装与配置", "30"),
        ("7.3", "启动机器人", "31"),
        ("7.4", "日常操作", "32"),
        ("7.5", "常见问题", "33"),
        ("8.", "开发者指南", "34"),
        ("8.1", "项目结构一览", "34"),
        ("8.2", "添加新的发送后端", "35"),
        ("8.3", "扩展联系人图谱", "35"),
        ("8.4", "调试技巧", "36"),
    ]

    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        indent = "    " if num.startswith("2.") or num.startswith("3.") or num.startswith("4.") or num.startswith("5.") or num.startswith("6.") or num.startswith("7.") or num.startswith("8.") else ""
        if not num[0].isdigit() or num.startswith("2.") or num.startswith("3.") or num.startswith("4.") or num.startswith("5.") or num.startswith("6.") or num.startswith("7.") or num.startswith("8."):
            indent = "    "
        else:
            indent = ""
        p.text = f"{indent}{num}  {title_text}"
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            if not num.startswith("2.") and not num.startswith("3.") and not num.startswith("4.") and not num.startswith("5.") and not num.startswith("6.") and not num.startswith("7.") and not num.startswith("8.") and num[0].isdigit():
                run.font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("1. 项目概述", level=1)

    doc.add_paragraph(
        "赵有才微信机器人是一个基于 Claude API 的智能微信自动回复系统。"
        "它以「赵有才」这一虚拟人设为核心，通过分析用户54,047条真实微信聊天记录构建人格模型，"
        "能根据不同的聊天对象（女朋友、铁哥们、学弟学妹、老师等）自动切换语气和回复风格，"
        "实现高度个性化的自动回复。"
    )

    doc.add_heading("核心特性", level=2)

    features = [
        ("多模态消息检测", "红点像素扫描 + 时间戳OCR扫描双重检测，自动发现未读消息"),
        ("人设感知回复", "基于60+联系人关系图谱，针对不同关系自动调整语气、篇幅、用词"),
        ("多后端发送", "支持剪贴板UI自动化（微信4.x）和WeChatFerry进程注入（微信3.9.5.81）两种发送模式"),
        ("Electron集成", "通过WeFlow桌面应用实现数据库级消息读取、可视化日志、手动回复发送"),
        ("OCR双引擎", "RapidOCR（PP-OCR中文模型）为主，Windows系统OCR为备，支持中文高精度识别"),
        ("自然延迟", "随机化回复延迟、模拟人类阅读时间，规避反自动化检测")
    ]

    for title_text, desc in features:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}：")
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(desc).font.size = Pt(10.5)

    doc.add_heading("运行模式一览", level=2)

    add_table_with_style(doc,
        ["模式", "启动方式", "消息检测", "消息发送", "适用场景"],
        [
            ["截图OCR模式\n（默认）", "python main.py", "截图+红点+时间戳OCR", "剪贴板粘贴", "微信4.x 任意版本"],
            ["WeChatFerry模式", "python main.py --wcf", "进程Hook直接取消息", "进程注入send_text", "微信3.9.5.81"],
            ["WCDB集成模式", "WeFlow Electron启动", "数据库轮询新消息", "可插拔后端(auto/wcf/clipboard)", "有WeFlow图形界面"],
        ],
        [3, 3.5, 3.5, 4, 3.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("2. 系统架构总览", level=1)

    doc.add_heading("2.1 整体数据流", level=2)

    doc.add_paragraph(
        "系统由三个层次组成：Electron前端（WeFlow桌面应用）、Python后端（AI引擎）、"
        "微信客户端本体。三者通过WebSocket和进程间通信（IPC）协作："
    )

    # Architecture table
    add_table_with_style(doc,
        ["层次", "组件", "职责", "通信方式"],
        [
            ["展示层", "BotPage.tsx / ChatPage.tsx / BotReplyBar.tsx",
             "UI展示、用户操作、事件日志", "IPC (contextBridge)"],
            ["桥接层", "Electron main.ts / botBridgeService.ts",
             "进程管理、WCDB数据库轮询、WebSocket客户端", "WebSocket → Python\nIPC ↔ Renderer"],
            ["AI引擎层", "bot_reply_server.py / bot_image.py",
             "消息处理、AI调用、回复生成", "WebSocket ← Electron\nClaude API → 云端"],
            ["发送层", "send_backend.py → wechat_send.py / wcferry",
             "消息投递（剪贴板或进程注入）", "Windows API / DLL注入"],
            ["数据层", "WeChat WCDB (SQLCipher)",
             "加密消息数据库（只读访问）", "wcdb_api.dll FFI"],
        ],
        [2.5, 4.5, 4.5, 4.0]
    )

    doc.add_heading("2.2 三大运行模式", level=2)

    doc.add_heading("模式一：截图OCR模式（默认 — main.py）", level=3)
    p = doc.add_paragraph(
        "本模式不依赖任何数据库访问，完全通过屏幕截图和OCR识别来工作。"
        "适用于所有微信版本（包括4.x），无需特殊权限。"
    )
    add_code_block(doc,
        "python main.py                    # 标准自动回复模式\n"
        "python main.py --quiet            # 静默模式，减少输出\n"
        "python main.py --select           # 手动标定聊天区域\n"
        "python main.py --once 张三        # 单次：读取+回复张三的消息后退出"
    )

    doc.add_heading("模式二：WeChatFerry模式（--wcf）", level=3)
    p = doc.add_paragraph(
        "通过wcferry库注入spy.dll到微信进程，直接Hook微信内部的SendMsg函数。"
        "零UI自动化、零截图、零OCR。仅支持微信3.9.5.81版本。"
    )
    add_code_block(doc,
        "python main.py --wcf              # 启动WeChatFerry后端\n"
        "# 前提：微信3.9.5.81已登录，pip install wcferry>=39.0.0"
    )

    doc.add_heading("模式三：WCDB集成模式（WeFlow Electron）", level=3)
    p = doc.add_paragraph(
        "由WeFlow桌面应用的BotBridgeService管理整个生命周期。"
        "通过wcdb_api.dll直接读取微信加密数据库获取新消息，"
        "通过WebSocket推送给Python AI引擎处理，回复通过可插拔发送后端投递。"
    )

    doc.add_heading("2.3 技术栈总览", level=2)

    add_table_with_style(doc,
        ["技术", "用途", "位置"],
        [
            ["Python 3.11", "后端AI引擎主语言", "wechat-bot/*.py"],
            ["TypeScript 5.x", "Electron前端", "WeFlow/src/"],
            ["React 18 + Zustand", "前端UI框架 + 状态管理", "WeFlow/src/"],
            ["Electron 30+", "桌面应用框架", "WeFlow/electron/"],
            ["Claude API (Anthropic)", "AI对话生成", "claude_client.py"],
            ["DeepSeek API", "Claude API代理（降低成本）", "config.py"],
            ["RapidOCR (PP-OCR ONNX)", "中文OCR引擎", "ocr_engine.py"],
            ["Windows OCR (winrt)", "OCR备选方案", "ocr_engine.py"],
            ["wcdb_api.dll (C++ FFI)", "微信加密数据库解密读取", "wcdbCore.ts"],
            ["wcferry (C++ DLL注入)", "微信进程Hook消息收发", "bot_wcf.py"],
            ["pyautogui + pyperclip", "UI自动化（剪贴板粘贴）", "wechat_send.py"],
            ["win32gui (pywin32)", "Windows窗口管理API", "wechat_send.py"],
            ["koffi (Node.js FFI)", "调用wcdb_api.dll的FFI库", "wcdbCore.ts"],
            ["WebSocket (RFC 6455)", "Python ↔ Electron 实时通信", "ws_server.py"],
        ],
        [4.0, 5.5, 5.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 3. PYTHON BACKEND MODULES
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("3. Python 后端模块详解", level=1)

    # 3.1 main.py
    doc.add_heading("3.1 入口调度 — main.py", level=2)
    doc.add_paragraph(
        "main.py 是整个Python后端的统一入口，负责解析命令行参数并分发给对应的后端实现。"
        "支持四种运行模式：默认截图OCR模式(--image)、传统wxauto模式(--legacy)、"
        "WeChatFerry模式(--wcf)、单次模式(--once NAME)。"
        "所有模式在启动前都会验证 ANTHROPIC_API_KEY 环境变量是否已设置。"
    )
    doc.add_paragraph("架构设计要点：", style='List Bullet')
    doc.add_paragraph("使用 argparse 提供清晰的命令行接口", style='List Bullet')
    doc.add_paragraph("每种模式对应独立的 Bot 类，避免代码耦合", style='List Bullet')
    doc.add_paragraph("--ws-port 参数允许在WCDB集成模式下指定WebSocket端口", style='List Bullet')
    doc.add_paragraph("--select 参数触发Tkinter区域选择器，用于手动标定截图区域", style='List Bullet')

    # 3.2 bot_image.py
    doc.add_heading("3.2 核心引擎 — bot_image.py（约1600行）", level=2)
    doc.add_paragraph(
        "这是整个项目最核心的模块，实现了完整的 截图→OCR→AI分析→发送 全链路。"
        "设计目标是完全独立于微信版本——不读数据库，不注入进程，"
        "只通过屏幕截图和Windows API操作微信窗口。"
    )

    doc.add_heading("核心技术要点", level=3)
    doc.add_paragraph("BitBlt窗口DC截图：微信4.x使用Qt5的DirectComposition渲染，常规的ImageGrab.grab()无法捕获窗口内容。"
                      "必须使用win32gui.GetWindowDC()获取窗口的设备上下文，再用BitBlt从窗口DC拷贝像素。", style='List Bullet')
    doc.add_paragraph("双模式未读检测：同时运行红点像素扫描（检测红色badge的RGB值）和时间戳OCR扫描（识别\"刚刚\"、时间变化），"
                      "合并结果后在30px距离内去重。红点检测覆盖率约70%（依赖微信主题），时间戳检测覆盖剩余30%。", style='List Bullet')
    doc.add_paragraph("OCR预处理管线：图像先放大4倍（如原图<700px），再进行2%-98%百分位对比度拉伸，最后用Unsharp Mask锐化。"
                      "这对提高中文OCR识别率至关重要。", style='List Bullet')
    doc.add_paragraph("布局自动标定：扫描窗口40%-55%水平区域，用3px宽条带逐列分析亮度、方差、暗像素数，"
                      "找到聊天列表和对话面板之间的分割线位置。支持手动覆盖（--select + DIVIDER_X配置）。", style='List Bullet')
    doc.add_paragraph("三重去重机制：精确指纹匹配（SHA-256）、子串匹配（长度>80%时）、"
                      "CJK字符集Jaccard相似度（阈值0.55），确保OCR识别偏差不会导致重复回复。", style='List Bullet')

    doc.add_heading("主循环流程", level=3)
    add_code_block(doc,
        "run_forever():\n"
        "  while True:\n"
        "    1. _detect_unread_chats()       # 红点扫描（右侧25%区域）\n"
        "    2. _scan_chat_list_activity()   # 时间戳OCR扫描（右侧30%区域）\n"
        "    3. _merge_nearby_positions()    # 合并+去重（30px阈值）\n"
        "    for each hit:\n"
        "      a. click_chat(hit)            # 点击进入聊天\n"
        "      b. read_last_messages()      # 截图+OCR读取消息\n"
        "      c. _extract_incoming_messages() # 解析OCR输出\n"
        "      d. _should_skip_message()    # 三重去重检查\n"
        "      e. build_prompt()            # 构建AI上下文\n"
        "      f. get_reply()               # 调用Claude API\n"
        "      g. type_and_send()           # 剪贴板粘贴发送\n"
        "      h. contact_prime()           # 防重复回复守卫\n"
        "    4. _deselect_chat()             # 取消选中聊天\n"
        "    5. periodic_cleanup()           # 定期清理过期状态"
    )

    # 3.3 bot_reply_server.py
    doc.add_heading("3.3 轻量引擎 — bot_reply_server.py", level=2)
    doc.add_paragraph(
        "专为WeFlow Electron集成设计的轻量级回复引擎。与bot_image.py不同，"
        "它不做任何截图、OCR、窗口检测——完全依赖WeFlow通过WebSocket推送消息。"
        "消息来自WeFlow通过wcdb_api.dll直接读取微信加密数据库，因此消息内容精确无误。"
    )
    doc.add_paragraph("架构特点：", style='List Bullet')
    doc.add_paragraph("无截图、无OCR、无窗口检测——完全无头(headless)运行", style='List Bullet')
    doc.add_paragraph("通过WebSocket接收命令（stop/pause/resume/reply/send_message）", style='List Bullet')
    doc.add_paragraph("使用可插拔发送后端(send_backend)，自动选择最佳发送方式", style='List Bullet')
    doc.add_paragraph("支持 --backend auto|wcf|clipboard 三种模式切换", style='List Bullet')

    # 3.4 bot_wcf.py
    doc.add_heading("3.4 WeChatFerry 后端 — bot_wcf.py", level=2)
    doc.add_paragraph(
        "利用wcferry库实现的零UI自动化后端。通过spy.dll注入微信进程，"
        "直接Hook微信内部的SendMsg函数来发送消息。"
        "消息接收也通过微信进程内Hook获取，无需截图和OCR。"
    )
    p = doc.add_paragraph()
    run = p.add_run("版本限制：")
    run.bold = True
    p.add_run(
        "仅支持微信3.9.5.81。微信4.x完全重构了内部架构（Qt5渲染、wcdb加密数据库、PPL任务模式），"
        "spy.dll中硬编码的函数偏移量在4.x中不存在，强行注入会导致微信崩溃。"
        "截至2026年6月，没有已知的社区分支支持微信4.x。"
    )

    # 3.5 send_backend.py
    doc.add_heading("3.5 统一发送模块 — send_backend.py", level=2)
    doc.add_paragraph(
        "最新引入的可插拔发送后端架构。定义了SendBackend抽象协议，"
        "提供ClipboardBackend（剪贴板UI自动化）和WcfBackend（WeChatFerry进程注入）两种实现。"
        "通过create_backend()工厂函数自动选择可用的最佳后端。"
    )
    doc.add_paragraph("设计模式：策略模式 + 工厂模式", style='List Bullet')
    doc.add_paragraph("WcfBackend的注册表预检：在尝试创建Wcf实例前，先检查Windows注册表中是否存在微信3.9.x的安装路径", style='List Bullet')
    doc.add_paragraph("ClipboardBackend：封装wechat_send.py的send_clipboard()，每次发送前调用release_modifiers()防止Ctrl键卡住", style='List Bullet')

    add_code_block(doc,
        "class SendBackend(Protocol):\n"
        "    name: str\n"
        "    def is_available(self) -> bool: ...\n"
        "    def send(self, text: str, contact: str | None = None) -> bool: ...\n"
        "\n"
        "create_backend(prefer='auto') -> SendBackend\n"
        "  # 'auto': 优先WcfBackend，不可用时回退ClipboardBackend\n"
        "  # 'wcf':  强制WcfBackend，不可用时回退\n"
        "  # 'clipboard': 强制ClipboardBackend"
    )

    # 3.6 wechat_send.py
    doc.add_heading("3.6 窗口操作模块 — wechat_send.py", level=2)
    doc.add_paragraph(
        "提供两层API：Tier 1是窗口管理的基础构建块（查找、激活、点击、清除），"
        "Tier 2是完整的发送函数（剪贴板模式和模拟打字模式）。"
        "所有发送都遵循6步协议：查找窗口→激活→点击输入区→清除→填入文本→按Enter。"
    )
    doc.add_paragraph("关键技术细节：", style='List Bullet')
    doc.add_paragraph("activate_window()使用Alt键技巧获取前台激活权限——先发送Alt按下/释放事件，"
                      "让Windows认为进程有用户输入，然后才能成功调用SetForegroundWindow()", style='List Bullet')
    doc.add_paragraph("click_input_area()在窗口底部往上130px处点击（微信4.x文本输入区的中间位置），"
                      "避免点击到底部50px处的工具栏按钮（会误触发截图）", style='List Bullet')
    doc.add_paragraph("navigate_to_chat()通过点击聊天列表最上方行（left+15%宽度, top+100px）来导航，"
                      "因为最新消息的发送者总在列表顶部", style='List Bullet')
    doc.add_paragraph("release_modifiers()在每次发送前释放LCtrl/RCtrl/LAlt/RAlt的KEYUP事件，"
                      "解决WeFlow中Ctrl+Enter快捷键导致Ctrl键卡住的问题", style='List Bullet')

    # 3.7 human_send.py
    doc.add_heading("3.7 人类化输入 — human_send.py", level=2)
    doc.add_paragraph(
        "通过Windows SendInput API的KEYEVENTF_UNICODE标志直接发送Unicode字符，"
        "模拟真人的打字节奏。支持变速输入、句子间停顿、爆冲模式（15%概率连续快速输入2-5个字符）、"
        "思考停顿等。"
    )
    doc.add_paragraph("延迟参数：CJK字符50-180ms/字，ASCII字符25-100ms/字，"
                      "句末停顿350-850ms，段落停顿700-1600ms。", style='List Bullet')
    doc.add_paragraph("注意：微信4.x的Qt5输入框对SendInput UNICODE支持不完善，"
                      "因此wechat_send.py已统一改用剪贴板粘贴方式。"
                      "本模块保留用于其他可能的应用场景和剪贴板的参考实现。", style='List Bullet')

    # 3.8 claude_client.py
    doc.add_heading("3.8 AI 客户端 — claude_client.py", level=2)
    doc.add_paragraph(
        "封装Anthropic Python SDK，提供统一的AI调用接口。被所有后端（bot_image.py、"
        "bot_wcf.py、bot_reply_server.py）共享使用。"
    )
    doc.add_paragraph("核心方法：", style='List Bullet')
    doc.add_paragraph("build_prompt(contact, incoming_text)：构建基础prompt，包含800字符预算内的对话历史", style='List Bullet')
    doc.add_paragraph("get_reply(messages, system_prompt)：调用Claude API，temperature=0.85，最大500 tokens。"
                      "自动去除回复两端的引号（AI有时会在回复外加引号）", style='List Bullet')
    doc.add_paragraph("注意：实际API调用可能通过ANTHROPIC_BASE_URL代理到DeepSeek以降低成本。"
                      "config.py中默认model为deepseek-chat。", style='List Bullet')

    # 3.9 conversation_manager.py
    doc.add_heading("3.9 对话管理 — conversation_manager.py", level=2)
    doc.add_paragraph(
        "管理每个联系人的对话历史。使用collections.deque实现最大长度限制（默认20条），"
        "自动淘汰旧消息。同时跟踪最后回复时间（用于冷却检查）和回复计数。"
    )
    doc.add_paragraph("normalize_contact()函数专门处理OCR导致的联系人名称畸变，"
                      "通过去除OCR噪点字符、折叠空白、优先保留ASCII词等方式在同一个人名的不同OCR结果之间进行匹配。", style='List Bullet')

    # 3.10 contacts.py
    doc.add_heading("3.10 联系人图谱 — contacts.py", level=2)
    doc.add_paragraph(
        "构建了约60个联系人的关系拓扑图，是赵有才「人设感知」回复的核心。"
        "每个联系人包含：沟通模式(mode)、双方角色(role)、关系动态(dynamic)、"
        "称呼列表(aliases，用于OCR匹配)、语气量级(tone，0-10)。"
    )
    p = doc.add_paragraph()
    run = p.add_run("联系人分类统计：")
    run.bold = True
    add_table_with_style(doc,
        ["类别", "人数", "典型沟通模式"],
        [
            ["铁哥们", "7", "完全放松（说脏话、互损、不装）"],
            ["女朋友", "1", "亲密（亲昵、撒娇、陪伴感）"],
            ["亲姐/家人", "4", "家人（随意但尊重、可求助）"],
            ["专业搭档", "3", "高效协作（技术讨论为主、简洁）"],
            ["学弟学妹", "16", "靠谱学长（耐心、鼓励、不过度社交）"],
            ["老师", "8", "尊重礼貌（正式、简洁、有分寸）"],
            ["其他朋友/同学", "13", "看人下菜碟（根据熟悉度调整）"],
        ],
    )

    # 3.11 zhaoyoucai_prompt.py
    doc.add_heading("3.11 人设提示词 — zhaoyoucai_prompt.py", level=2)
    doc.add_paragraph(
        "约5000词的系统提示词（System Prompt），是整个AI回复系统的灵魂。"
        "基于赵有才真实的54,047条微信消息、281个联系人的聊天数据提炼而成。"
    )
    doc.add_paragraph("提示词结构：", style='List Bullet')
    doc.add_paragraph("第零层：硬规则 — PII保护、越狱检测、始终入戏、非心理咨询师、隐私保护", style='List Bullet')
    doc.add_paragraph("信息分级L0-L4：定义不同级别信息对不同关系的可见性", style='List Bullet')
    doc.add_paragraph("身份名片：年龄、学校、专业、家庭、女友、副业、游戏、运动", style='List Bullet')
    doc.add_paragraph("语言风格：基于数据分析的说话习惯——消息长度分布、「啥」vs「什么」的词频比（4:6）、"
                      "[捂脸]表情使用统计（共6557次）、道歉模式、冲突解决风格", style='List Bullet')
    doc.add_paragraph("关系语言切换：6种不同沟通模式的语言特点", style='List Bullet')
    doc.add_paragraph("决策与情绪模式：理性决策者、三级愤怒递进、行动型浪漫", style='List Bullet')

    # 3.12 ocr_engine.py
    doc.add_heading("3.12 OCR 引擎 — ocr_engine.py", level=2)
    doc.add_paragraph(
        "双引擎OCR抽象层。主引擎为RapidOCR（PP-OCR中文模型通过ONNX Runtime推理），"
        "备选引擎为Windows内置OCR（通过winrt调用）。"
    )
    doc.add_paragraph("预处理管线：", style='List Bullet')
    doc.add_paragraph("小图放大：宽度<700px时放大4倍（Lanczos重采样）", style='List Bullet')
    doc.add_paragraph("对比度拉伸：2nd-98th百分位值映射到0-255", style='List Bullet')
    doc.add_paragraph("Unsharp Mask锐化：radius=1.5, amount=150%，补偿放大导致的模糊", style='List Bullet')
    doc.add_paragraph("CJK空格折叠：移除中文字符之间Windows OCR引入的多余空格", style='List Bullet')

    # 3.13 ws_server.py
    doc.add_heading("3.13 WebSocket 服务 — ws_server.py", level=2)
    doc.add_paragraph(
        "纯Python标准库实现的WebSocket服务器（RFC 6455），无任何外部依赖。"
        "运行在独立守护线程中，支持广播推送和命令接收双向通信。"
    )
    doc.add_paragraph("实现细节：", style='List Bullet')
    doc.add_paragraph("基于asyncio的TCP服务器，手动实现HTTP Upgrade握手", style='List Bullet')
    doc.add_paragraph("仅支持文本帧（opcode 0x1），处理close/ping控制帧", style='List Bullet')
    doc.add_paragraph("线程安全：broadcast_sync()通过asyncio.run_coroutine_threadsafe()从主线程安全推送", style='List Bullet')
    doc.add_paragraph("命令队列：pop_commands()由bot主循环轮询，线程锁保护", style='List Bullet')
    doc.add_paragraph("心跳机制：每10秒向所有客户端广播heartbeat事件", style='List Bullet')

    # 3.14 display.py
    doc.add_heading("3.14 终端显示 — display.py", level=2)
    doc.add_paragraph(
        "为截图OCR模式提供的ANSI终端可视化界面。使用stderr输出持久状态行"
        "（使用\\r覆盖更新），stdout输出滚动事件日志。"
        "通过颜色编码区分不同事件类型（收到消息=青色、已发送回复=绿色、错误=红色）。"
    )

    # 3.15 config.py
    doc.add_heading("3.15 配置中心 — config.py", level=2)
    doc.add_paragraph("所有可调参数集中在一个文件中：")

    add_table_with_style(doc,
        ["参数", "默认值", "说明"],
        [
            ["ANTHROPIC_API_KEY", "(环境变量)", "Anthropic API密钥（必填）"],
            ["ANTHROPIC_MODEL", "deepseek-chat", "使用的AI模型"],
            ["MAX_CONTEXT_MESSAGES", "20", "对话历史最大条数"],
            ["MAX_TOKENS", "500", "AI回复最大token数"],
            ["REPLY_DELAY_MIN / MAX", "1.0 / 4.0", "回复前随机延迟范围（秒）"],
            ["POLL_INTERVAL", "2.0", "轮询间隔（秒）"],
            ["REPLY_COOLDOWN", "12.0", "同一联系人冷却时间（秒）"],
            ["DIVIDER_X", "None", "聊天列表分割线手动位置"],
            ["DEBUG", "True", "调试模式开关"],
        ],
        [4.5, 3.5, 7.0]
    )

    # 3.16 auxiliary
    doc.add_heading("3.16 辅助工具模块", level=2)

    doc.add_paragraph("diagnose.py：独立诊断工具，捕获微信窗口，分析列布局用于分割线检测，"
                      "测试不同分割线位置的OCR效果，检测红点，保存标注图像。用于调试布局问题。", style='List Bullet')
    doc.add_paragraph("region_selector.py：Tkinter全屏透明覆盖层，允许用户手动拖拽绘制矩形来标定聊天列表和对话面板区域。"
                      "通过--select参数触发。支持DPI感知坐标系。", style='List Bullet')
    doc.add_paragraph("test_ocr.py：RapidOCR vs Windows OCR对比测试工具，并排显示两个引擎的识别结果和耗时，"
                      "用于评估OCR质量。", style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 4. ELECTRON FRONTEND
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("4. Electron 前端详解", level=1)

    doc.add_heading("4.1 主进程 — main.ts", level=2)
    doc.add_paragraph(
        "WeFlow的Electron主进程（约4500行），负责窗口管理、自动更新、配置管理、"
        "以及整个应用的IPC处理器注册。Bot相关的IPC处理器约15个，"
        "全部路由到BotBridgeService单例。"
    )
    doc.add_paragraph("Bot相关IPC处理器列表：", style='List Bullet')
    add_code_block(doc,
        "ipcMain.handle('bot:start', ...)          # 启动Python机器人进程\n"
        "ipcMain.handle('bot:stop', ...)           # 停止机器人\n"
        "ipcMain.handle('bot:pause', ...)          # 暂停回复\n"
        "ipcMain.handle('bot:resume', ...)         # 恢复回复\n"
        "ipcMain.handle('bot:getStatus', ...)      # 获取运行状态\n"
        "ipcMain.handle('bot:getConfig', ...)      # 获取配置\n"
        "ipcMain.handle('bot:setConfig', ...)      # 修改配置\n"
        "ipcMain.handle('bot:getStats', ...)       # 获取统计数据\n"
        "ipcMain.handle('bot:getLogs', ...)        # 获取事件日志\n"
        "ipcMain.handle('bot:clearLogs', ...)      # 清空日志\n"
        "ipcMain.handle('bot:sendMessage', ...)    # 手动发送消息\n"
        "ipcMain.handle('bot:testConnection', ...) # WebSocket连通性测试"
    )

    doc.add_heading("4.2 桥接服务 — botBridgeService.ts", level=2)
    doc.add_paragraph(
        "BotBridgeService是Electron和Python之间的核心桥梁，管理Python进程的整个生命周期。"
        "它是一个单例服务，负责进程创建、WebSocket连接、WCDB数据库轮询、事件转发。"
    )
    doc.add_paragraph("核心工作流程：", style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("启动流程：")
    run.bold = True
    add_code_block(doc,
        "start():\n"
        "  1. 确保WCDB服务已连接\n"
        "  2. 查找可用端口（首选9877，自动回退）\n"
        "  3. spawn('python bot_reply_server.py --ws-port PORT')\n"
        "  4. 等待1.5秒让Python初始化\n"
        "  5. 连接WebSocket到Python\n"
        "  6. 启动WCDB轮询（每2秒检查新消息）\n"
        "  7. 转发状态到Renderer"
    )

    p = doc.add_paragraph()
    run = p.add_run("消息检测机制：")
    run.bold = True
    doc.add_paragraph(
        "botBridgeService通过wcdbService读取微信加密数据库。启动时先「播种」——"
        "将当前所有会话的最新3条消息标记为已处理。之后每2秒轮询：遍历所有会话，"
        "检查每个会话的最新消息，如果发现未在processedMessages集合中的"
        "新消息（isSend===0表示收到、非自己的消息），则通过WebSocket推送给Python。"
        "processedMessages集合上限10000条，防止内存泄漏。"
    )

    p = doc.add_paragraph()
    run = p.add_run("停止流程：")
    run.bold = True
    add_code_block(doc,
        "stop():\n"
        "  1. 停止WCDB轮询定时器\n"
        "  2. 发送 {command:'stop'} 到Python\n"
        "  3. 等待800ms 让Python优雅退出\n"
        "  4. 发送SIGTERM信号\n"
        "  5. 3秒后未退出则发送SIGKILL强制终止\n"
        "  6. 断开WebSocket连接"
    )

    doc.add_heading("4.3 预加载脚本 — preload.ts", level=2)
    doc.add_paragraph(
        "使用Electron的contextBridge机制安全地暴露API给渲染进程。"
        "Bot相关API封装为window.electronAPI.bot对象，"
        "提供start/stop/pause/resume/getStatus/getConfig/setConfig/getStats/"
        "getLogs/clearLogs/sendMessage/testConnection等方法，"
        "以及onEvent/onStatusChange事件订阅。每条方法调用ipcRenderer.invoke，"
        "事件通过ipcRenderer.on监听主进程推送。"
    )

    doc.add_heading("4.4 Bot管理页 — BotPage.tsx", level=2)
    doc.add_paragraph(
        "机器人管理UI，包含三个选项卡："
    )
    doc.add_paragraph("事件日志（Terminal图标）：实时滚动显示机器人所有事件——收到消息（蓝）、"
                      "已发送回复（绿）、扫描（灰）、错误（红）、心跳（浅灰）、状态变更（蓝）。"
                      "支持自动滚动和手动滚动。", style='List Bullet')
    doc.add_paragraph("配置（Settings图标）：API密钥（密码输入框，自动保存800ms防抖）、模型名称、"
                      "回复延迟范围、冷却时间、轮询/扫描间隔。有保存按钮和脏状态跟踪。", style='List Bullet')
    doc.add_paragraph("统计（Zap图标）：6个统计卡片（收到/发送/API调用/错误/扫描/运行时间）+"
                      "回复率百分比。", style='List Bullet')

    doc.add_heading("4.5 回复输入栏 — BotReplyBar.tsx", level=2)
    doc.add_paragraph(
        "嵌入在ChatPage每个好友对话框底部的消息发送栏，模仿微信原生输入框布局。"
        "支持Ctrl+Enter快捷键发送，2000字符上限，实时显示预估打字时间。"
        "机器人未连接时显示为禁用状态，提示「请先启动机器人」。"
    )
    doc.add_paragraph("关键实现细节：", style='List Bullet')
    doc.add_paragraph("Ctrl+Enter处理：同时调用preventDefault()和stopPropagation()，"
                      "然后blur()移除焦点，防止事件冒泡导致意外行为（如截图）", style='List Bullet')
    doc.add_paragraph("发送状态：sending → sent（绿色）/ error（红色），过程中Send按钮替换为旋转动画", style='List Bullet')
    doc.add_paragraph("消息通过window.electronAPI.bot.sendMessage(contactName, message)发送，"
                      "经IPC→WebSocket→Python后端链到达发送后端", style='List Bullet')

    doc.add_heading("4.6 状态管理 — botStore.ts", level=2)
    doc.add_paragraph(
        "基于Zustand的全局状态管理。存储机器人的运行状态(BotStatus)、配置(BotConfig)、"
        "统计数据(BotStats)、事件日志(BotEvent[]，上限500条)。"
        "addEvent()操作在添加事件的同时自动更新对应的统计计数器（消息数、回复数、扫描数等），"
        "保证UI和数据的一致性。"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 5. AI SYSTEM
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("5. AI 智能回复系统", level=1)

    doc.add_heading("5.1 多维度上下文构建", level=2)
    doc.add_paragraph(
        "每次AI调用前，系统会构建一个多层次的上下文提示词(messages数组)，"
        "包含以下维度："
    )

    add_table_with_style(doc,
        ["维度", "来源", "内容"],
        [
            ["系统人设", "zhaoyoucai_prompt.py", "约5000词的系统提示词，定义赵有才的完整人设"],
            ["对话历史", "conversation_manager.py", "最近20条消息，智能选择（1000字符预算）"],
            ["关系图谱", "contacts.py", "与当前联系人的关系模式、角色、互动动态"],
            ["新消息", "当前输入", "联系人刚发来的消息原文"],
            ["第三方提及", "contacts.py find_mentioned", "如果消息中提到其他已知联系人，附带相关背景"],
        ],
    )

    doc.add_paragraph(
        "上下文选择算法（_select_context）：始终保留最近4条消息，从第5条开始倒序添加，"
        "跳过长度≤2的琐碎消息（嗯、好、OK等），在1000字符预算内尽可能多包含历史。"
        "如果历史被截断，在开头插入\"…（更早的对话已省略）\"标记。"
    )

    doc.add_heading("5.2 人设系统设计", level=2)
    doc.add_paragraph(
        "赵有才的人设基于54,047条真实微信消息数据提炼。这不是一个模板化的bot——"
        "它的语言风格、用词习惯、甚至标点偏好都来自真实数据统计。"
    )

    doc.add_paragraph("数据驱动的语言特征：", style='List Bullet')
    doc.add_paragraph("\"啥\" vs \"什么\"的使用比例约为4:6，根据聊天对象自动在两者间切换", style='List Bullet')
    doc.add_paragraph("[捂脸]是使用最多的表情（共6557次），在不同情境下有自然的出现频率", style='List Bullet')
    doc.add_paragraph("消息长度分布因对象而异：对女朋友平均20+字/条，对铁哥们5-15字/条，对老师3-10字/条", style='List Bullet')

    doc.add_paragraph("6种沟通模式：", style='List Bullet')

    add_table_with_style(doc,
        ["模式", "适用对象", "语言特征"],
        [
            ["完全放松", "铁哥们", "脏话、互损、简洁、口语化、不用敬语"],
            ["亲密", "女朋友", "亲昵称呼、撒娇、陪伴感、主动分享、浪漫但行动型"],
            ["极简", "不熟的人", "简短回复、礼貌有距离、不主动找话题"],
            ["靠谱学长", "学弟学妹", "耐心鼓励、技术帮助、不过度社交、有事说事"],
            ["高效协作", "专业搭档", "技术讨论为主、简洁直接、不闲聊"],
            ["尊重礼貌", "老师/长辈", "正式用词、有分寸、简洁、礼貌谦逊"],
        ],
    )

    doc.add_heading("5.3 对话历史管理", level=2)
    doc.add_paragraph(
        "每个联系人独立维护对话历史（conversation_manager.py）。"
        "历史使用deque(maxlen=20)实现，超出容量自动淘汰最早的消息。"
        "同时跟踪last_reply_time（用于12秒冷却期检查，防止对同一人连续轰炸回复）"
        "和last_activity（用于1小时过期清理）。"
    )

    doc.add_heading("5.4 API 调用策略", level=2)
    doc.add_paragraph("模型配置：默认使用deepseek-chat（通过ANTHROPIC_BASE_URL代理），"
                      "也支持直接调用Anthropic Claude模型", style='List Bullet')
    doc.add_paragraph("参数设置：temperature=0.85（保证一定创造性但不过于随机），"
                      "max_tokens=500（每条回复约250-500字）", style='List Bullet')
    doc.add_paragraph("后处理：自动去除回复两端可能存在的引号（AI有时会将回复放在引号中）", style='List Bullet')
    doc.add_paragraph("错误处理：API调用失败时返回None，上层记录错误但不中断运行", style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 6. MESSAGE SENDING
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("6. 消息发送机制", level=1)

    doc.add_paragraph(
        "消息发送是整个系统与微信交互的最终环节。根据运行环境和微信版本，"
        "系统提供两种发送后端：ClipboardBackend（剪贴板UI自动化）和"
        "WcfBackend（WeChatFerry进程注入）。"
    )

    doc.add_heading("6.1 ClipboardBackend — 剪贴板模式", level=2)
    doc.add_paragraph("适用场景：微信4.x（及所有版本），需要微信窗口可见", style='List Bullet')
    doc.add_paragraph("工作原理：", style='List Bullet')

    add_code_block(doc,
        "send_clipboard(text, contact):\n"
        "  1. find_wechat_window()        # 枚举窗口找微信\n"
        "  2. activate_window(hwnd)       # 恢复+前置窗口（Alt键技巧）\n"
        "  3. release_modifiers()         # 释放卡住的Ctrl/Alt键\n"
        "  4. navigate_to_chat(contact)   # 点击聊天列表顶行\n"
        "  5. click_input_area(hwnd)      # 点击输入区域(bottom-130px)\n"
        "  6. clear_input()               # Ctrl+A + Backspace\n"
        "  7. pyperclip.copy(text)        # 复制到剪贴板\n"
        "  8. pyautogui.hotkey('ctrl','v')# 粘贴\n"
        "  9. pyautogui.press('enter')    # 发送"
    )

    doc.add_paragraph("关键注意事项：", style='List Bullet')
    doc.add_paragraph("activate_window需要Alt键技巧：Windows安全机制禁止后台进程调用SetForegroundWindow。"
                      "通过先模拟Alt键按下/释放，让系统认为进程有用户输入，从而获得前台激活权限。", style='List Bullet')
    doc.add_paragraph("click_input_area的坐标计算：微信4.x窗口底部30-80px是工具栏（含截图按钮），"
                      "如果在bottom-50处点击会触发截图。正确位置是bottom-130，即文本输入区域的中间。", style='List Bullet')
    doc.add_paragraph("release_modifiers是WeFlow集成的关键：用户在BotReplyBar按Ctrl+Enter发送时，"
                      "Ctrl键可能仍被物理按住。Python在操作微信前必须发送KEYUP释放所有修饰键，"
                      "否则后续的Ctrl+V粘贴会变成其他快捷键。", style='List Bullet')

    doc.add_heading("6.2 WcfBackend — WeChatFerry 模式", level=2)
    doc.add_paragraph("适用场景：微信3.9.5.81，微信可以最小化或隐藏", style='List Bullet')
    doc.add_paragraph("工作原理：", style='List Bullet')

    add_code_block(doc,
        "WcfBackend.send(text, contact):\n"
        "  1. _resolve_wxid(contact)      # 显示名→wxid映射\n"
        "     - 精确匹配缓存中的remark/name\n"
        "     - 模糊匹配（子串包含）\n"
        "     - 以wxid_开头则直接使用\n"
        "  2. wcf.send_text(text, wxid)   # RPC调用发送\n"
        "     - 序列化为protobuf\n"
        "     - 通过nanomsg TCP发送到spy.dll\n"
        "     - spy.dll调用微信内部SendMsg\n"
        "  3. 返回0=成功, 非0=失败"
    )

    doc.add_paragraph("注册表预检：在尝试创建Wcf实例前，先检查HKCU\\Software\\Tencent\\WeChat注册表键是否存在。"
                      "微信4.x使用Software\\Tencent\\Weixin（不同键名），因此预检会失败并安全回退，"
                      "避免了wcferry C++ SDK找不到微信时调用exit()导致Python进程崩溃的问题。", style='List Bullet')

    doc.add_heading("6.3 发送流程对比", level=2)

    add_table_with_style(doc,
        ["特性", "ClipboardBackend", "WcfBackend"],
        [
            ["微信版本", "全部版本", "仅3.9.5.81"],
            ["微信窗口", "必须可见", "可隐藏/最小化"],
            ["发送延迟", "~800ms（含窗口操作）", "~10ms（RPC调用）"],
            ["焦点切换", "会切换到微信窗口", "无（后台发送）"],
            ["中文支持", "优秀（剪贴板粘贴）", "优秀（直接调用）"],
            ["可靠性", "中等（受窗口状态影响）", "高（进程级Hook）"],
            ["反检测", "模拟人类操作", "完全无UI痕迹"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 7. USER MANUAL
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("7. 用户使用手册", level=1)

    doc.add_heading("7.1 环境要求", level=2)

    add_table_with_style(doc,
        ["组件", "要求", "备注"],
        [
            ["操作系统", "Windows 10/11 (x64)", "不支持macOS/Linux"],
            ["Python", "3.11+", "安装到PATH"],
            ["微信PC版", "4.x 或 3.9.5.81", "3.9.5.81支持更稳定的WcfBackend"],
            ["Node.js", "18+", "用于运行WeFlow Electron应用"],
            ["Anthropic API Key", "或 DeepSeek API Key", "设置到环境变量"],
            ["WeFlow", "已编译的Electron应用", "提供图形界面和WCDB访问"],
        ],
    )

    doc.add_heading("7.2 安装与配置", level=2)

    doc.add_heading("步骤一：安装Python依赖", level=3)
    add_code_block(doc,
        "cd C:\\Users\\Trouvailler\\wechat-bot\n"
        "pip install -r requirements.txt"
    )

    doc.add_paragraph("requirements.txt 核心依赖：", style='List Bullet')
    add_code_block(doc,
        "pyautogui          # 鼠标点击和键盘模拟\n"
        "pyperclip          # 剪贴板操作\n"
        "pywin32            # Windows API (窗口管理)\n"
        "Pillow             # 图像处理\n"
        "numpy              # 像素级计算\n"
        "anthropic          # Claude API客户端\n"
        "rapidocr-onnxruntime  # OCR引擎（可选但推荐）\n"
        "wcferry>=39.0.0   # WeChatFerry支持（可选，需微信3.9.5.81）"
    )

    doc.add_heading("步骤二：配置API密钥", level=3)
    doc.add_paragraph("方式一：通过WeFlow界面配置（推荐）")
    doc.add_paragraph("  打开WeFlow → 智能回复 → 配置选项卡 → 输入API Key → 自动保存")
    doc.add_paragraph("方式二：设置系统环境变量")
    add_code_block(doc,
        "# PowerShell:\n"
        "$env:ANTHROPIC_API_KEY = \"your-api-key-here\"\n"
        "$env:ANTHROPIC_BASE_URL = \"https://api.deepseek.com\"  # 如果使用DeepSeek代理\n"
        "$env:ANTHROPIC_MODEL = \"deepseek-chat\"                # 使用的模型"
    )

    doc.add_heading("步骤三：验证安装", level=3)
    add_code_block(doc,
        "cd C:\\Users\\Trouvailler\\wechat-bot\n"
        "python main.py --test          # 测试AI API连通性\n"
        "python -c \"from wechat_send import find_wechat_window; print(find_wechat_window())\"\n"
        "                                # 测试微信窗口检测"
    )

    doc.add_heading("7.3 启动机器人", level=2)

    doc.add_heading("通过WeFlow图形界面（推荐日常使用）", level=3)
    doc.add_paragraph("1. 启动WeFlow应用", style='List Number')
    doc.add_paragraph("2. 左侧导航 → 智能回复（Bot图标）", style='List Number')
    doc.add_paragraph("3. 点击「启动」按钮", style='List Number')
    doc.add_paragraph("4. 观察日志确认「状态: running」和「WS已连接」", style='List Number')
    doc.add_paragraph("5. 在聊天功能中打开好友对话框，即可在底部Bar手动回复", style='List Number')

    doc.add_heading("通过命令行（高级用户/调试）", level=3)

    doc.add_paragraph("截图OCR模式（微信4.x）：", style='List Bullet')
    add_code_block(doc, "python main.py")

    doc.add_paragraph("WeChatFerry模式（微信3.9.5.81）：", style='List Bullet')
    add_code_block(doc, "python main.py --wcf")

    doc.add_paragraph("WCDB集成模式（配合WeFlow使用）：", style='List Bullet')
    add_code_block(doc, "python bot_reply_server.py --ws-port 9877 --backend auto")

    doc.add_heading("7.4 日常操作", level=2)

    doc.add_paragraph("暂停/恢复：在智能回复页面点击「暂停」按钮，机器人停止回复但保持运行。"
                      "点击「继续」恢复。可用于需要人工接管对话的场景。", style='List Bullet')
    doc.add_paragraph("手动回复：在聊天功能中打开联系人对话框，在底部输入框输入内容，"
                      "按Ctrl+Enter发送。发送状态会实时显示在输入框右侧。", style='List Bullet')
    doc.add_paragraph("查看统计：在智能回复→统计选项卡查看消息数、回复率、API调用次数等数据。", style='List Bullet')
    doc.add_paragraph("清空日志：日志超过500条后旧日志会自动清理，也可以手动点击「清空日志」。", style='List Bullet')

    doc.add_heading("7.5 常见问题", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Q: 自动回复发送失败，提示SetForegroundWindow错误？")
    run.bold = True
    doc.add_paragraph("A: 这是Windows安全限制。已通过Alt键技巧修复。如果仍然失败，"
                      "请确保微信窗口没有被最小化到系统托盘。如果使用微信3.9.5.81，"
                      "可以在WeFlow中设置--backend wcf使用无UI模式。")

    p = doc.add_paragraph()
    run = p.add_run("Q: 手动回复(Ctrl+Enter)时微信截图按钮被触发？")
    run.bold = True
    doc.add_paragraph("A: 这是因为Ctrl键卡住导致Ctrl+V变成了其他快捷键。"
                      "最新版本已在send_backend中每次发送前调用release_modifiers()释放所有修饰键。"
                      "如果问题仍然存在，请重启WeFlow和微信。")

    p = doc.add_paragraph()
    run = p.add_run("Q: 机器人回复后自动暂停？")
    run.bold = True
    doc.add_paragraph("A: 这是旧版本的bug（navigate_to_chat点击到错误位置导致发送失败，"
                      "异常未捕获导致进程崩溃）。最新版本已将send调用包裹在try/except中并记录错误日志。"
                      "请更新到最新代码。")

    p = doc.add_paragraph()
    run = p.add_run("Q: 如何让机器人在微信最小化时也能回复？")
    run.bold = True
    doc.add_paragraph("A: 需要安装微信3.9.5.81并使用WcfBackend（--backend wcf）。"
                      "微信4.x目前不支持后台发送——微信4.x完全重构了内部架构，"
                      "wcferry的spy.dll无法Hook其函数。这是已知的技术限制。")

    p = doc.add_paragraph()
    run = p.add_run("Q: OCR识别不准确导致回复了错误的人？")
    run.bold = True
    doc.add_paragraph("A: 使用WeFlow WCDB模式可以完全避免OCR识别问题，"
                      "因为消息直接从数据库读取，内容100%准确。"
                      "如果使用截图OCR模式，请确保微信窗口分辨率不低于1600x900，"
                      "并使用RapidOCR引擎（pip install rapidocr-onnxruntime）。")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 8. DEVELOPER GUIDE
    # ══════════════════════════════════════════════════════════════════

    doc.add_heading("8. 开发者指南", level=1)

    doc.add_heading("8.1 项目结构一览", level=2)

    add_table_with_style(doc,
        ["路径", "说明"],
        [
            ["wechat-bot/", "Python后端根目录"],
            ["  main.py", "统一CLI入口 + 后端调度"],
            ["  bot_image.py", "截图OCR模式核心引擎（~1600行）"],
            ["  bot_reply_server.py", "WCDB集成模式轻量引擎"],
            ["  bot_wcf.py", "WeChatFerry模式后端"],
            ["  bot.py", "wxauto4传统模式后端（已弃用）"],
            ["  claude_client.py", "Claude/DeepSeek API封装"],
            ["  conversation_manager.py", "对话历史管理（deque）"],
            ["  contacts.py", "60+联系人关系图谱"],
            ["  zhaoyoucai_prompt.py", "5000词系统人设提示词"],
            ["  send_backend.py", "可插拔发送后端（Clipboard+Wcf）"],
            ["  wechat_send.py", "微信窗口操作 + 剪贴板发送"],
            ["  human_send.py", "SendInput人类化打字引擎"],
            ["  ocr_engine.py", "RapidOCR + Windows OCR双引擎"],
            ["  display.py", "ANSI终端可视化界面"],
            ["  ws_server.py", "纯Python WebSocket服务器"],
            ["  config.py", "全局配置参数"],
            ["  diagnose.py", "布局诊断工具"],
            ["  region_selector.py", "Tkinter区域选择器"],
            ["  test_ocr.py", "OCR引擎对比测试"],
            ["", ""],
            ["WeFlow-main/", "Electron前端根目录"],
            ["  electron/main.ts", "Electron主进程 + IPC处理"],
            ["  electron/preload.ts", "contextBridge安全API暴露"],
            ["  electron/services/botBridgeService.ts", "Python进程管理+WCDB轮询"],
            ["  electron/services/wcdbCore.ts", "wcdb_api.dll FFI绑定（~4600行）"],
            ["  src/pages/BotPage.tsx", "机器人管理页面"],
            ["  src/pages/ChatPage.tsx", "聊天分析页面"],
            ["  src/components/BotReplyBar.tsx", "聊天页底部回复输入栏"],
            ["  src/stores/botStore.ts", "Zustand全局状态"],
        ],
    )

    doc.add_heading("8.2 添加新的发送后端", level=2)
    doc.add_paragraph("实现SendBackend协议，注册到create_backend工厂即可。示例代码：")

    add_code_block(doc,
        "# 在 send_backend.py 中添加新的后端类\n"
        "class MyNewBackend:\n"
        "    \"\"\"通过XXX方式发送微信消息\"\"\"\n"
        "\n"
        "    @property\n"
        "    def name(self) -> str:\n"
        "        return \"MyNewBackend\"\n"
        "\n"
        "    def is_available(self) -> bool:\n"
        "        # 检查前置条件是否满足\n"
        "        return check_xxx_installed()\n"
        "\n"
        "    def send(self, text: str, contact: str | None = None) -> bool:\n"
        "        # 实现具体的发送逻辑\n"
        "        return do_send(text, contact)\n"
        "\n"
        "# 在 create_backend() 工厂函数中添加新的分支\n"
        "def create_backend(prefer='auto'):\n"
        "    if prefer == 'mynew':\n"
        "        return MyNewBackend()\n"
        "    # ... 其余逻辑"
    )

    doc.add_heading("8.3 扩展联系人图谱", level=2)
    doc.add_paragraph("在contacts.py的CONTACTS字典中添加新联系人：")

    add_code_block(doc,
        "CONTACTS = {\n"
        "    # ... 现有联系人\n"
        "    \"新朋友\": {\n"
        "        \"mode\": \"靠谱学长\",       # 6种模式之一\n"
        "        \"role\": \"技术指导者\",      # 赵有才在这个关系中的角色\n"
        "        \"aliases\": [\"新朋友微信名\", \"别名2\"],  # OCR匹配用\n"
        "        \"dynamic\": (\n"
        "            \"1. 你们是大学同一个实验室的，他比你低2级\\n\"\n"
        "            \"2. 他经常问你Python相关的问题，你很乐意解答\\n\"\n"
        "            \"3. 保持学长风范，耐心但不过度社交\\n\"\n"
        "        ),\n"
        "        \"tone\": 5                  # 0=极简 10=极详细\n"
        "    },\n"
        "}"
    )

    doc.add_heading("8.4 调试技巧", level=2)

    doc.add_paragraph("1. 测试AI API连通性：python main.py --test", style='List Number')
    doc.add_paragraph("2. 诊断布局问题：python diagnose.py（生成标注图像）", style='List Number')
    doc.add_paragraph("3. 手动标定区域：python main.py --select（拖拽选择聊天区域）", style='List Number')
    doc.add_paragraph("4. 比较OCR引擎：python test_ocr.py debug_screenshot.png --compare", style='List Number')
    doc.add_paragraph("5. 单次回复测试：python main.py --once 张三（仅处理张三的消息一次）", style='List Number')
    doc.add_paragraph("6. 查看WebSocket通信：启动bot_reply_server.py后观察控制台日志", style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("日志位置：")
    run.bold = True
    doc.add_paragraph("Python端：控制台stdout（带时间戳和颜色编码）", style='List Bullet')
    doc.add_paragraph("Electron端：WeFlow → 智能回复 → 事件日志（可查看最近500条）", style='List Bullet')
    doc.add_paragraph("开发者工具：WeFlow按F12打开DevTools，在Console中查看IPC通信", style='List Bullet')

    # ══════════════════════════════════════════════════════════════════
    # APPENDIX
    # ══════════════════════════════════════════════════════════════════

    doc.add_page_break()
    doc.add_heading("附录", level=1)

    doc.add_heading("附录A：WebSocket 协议参考", level=2)

    add_table_with_style(doc,
        ["方向", "消息格式", "说明"],
        [
            ["Python→Electron", '{"type":"status","state":"running","uptime":0}', "状态更新"],
            ["Python→Electron", '{"type":"msg_received","contact":"张三","text":"..."}', "收到新消息"],
            ["Python→Electron", '{"type":"reply_sent","contact":"张三","text":"..."}', "已发送回复"],
            ["Python→Electron", '{"type":"error","message":"..."}', "错误信息"],
            ["Python→Electron", '{"type":"heartbeat","scan_count":42}', "心跳（每10秒）"],
            ["Electron→Python", '{"command":"reply","contact":"张三","message":"...","displayName":"..."}', "请求AI回复"],
            ["Electron→Python", '{"command":"send_message","contact":"张三","message":"..."}', "手动发送消息"],
            ["Electron→Python", '{"command":"stop"}', "停止机器人"],
            ["Electron→Python", '{"command":"pause"} / {"command":"resume"}', "暂停/恢复"],
        ],
    )

    doc.add_heading("附录B：微信4.x 窗口布局参考", level=2)

    doc.add_paragraph("以下布局数据基于微信4.x (Qt5.15.14) 在1920×1080分辨率下测得：")
    add_code_block(doc,
        "窗口区域（从顶部算）：\n"
        "  0-30px    标题栏\n"
        "  30-60px   搜索框\n"
        "  60px+     聊天列表\n"
        "\n"
        "窗口区域（从底部算）：\n"
        "  0-30px    窗口边框\n"
        "  30-80px   工具栏（表情、截图、文件等按钮）⚠ 避免点击此区域\n"
        "  80-220px  文本输入区域 ✅ 安全点击区\n"
        "\n"
        "聊天列表分割线：窗口宽度的40%-55%处（自动标定）\n"
        "新消息联系人：始终在聊天列表最顶部"
    )

    doc.add_heading("附录C：依赖完整列表", level=2)

    add_table_with_style(doc,
        ["包名", "版本要求", "用途"],
        [
            ["pyautogui", ">=0.9.50", "鼠标/键盘自动化"],
            ["pyperclip", ">=1.8.0", "剪贴板操作"],
            ["pywin32", ">=305", "Windows API"],
            ["Pillow", ">=10.0.0", "图像处理"],
            ["numpy", ">=1.24.0", "像素数组计算"],
            ["anthropic", ">=0.30.0", "Claude API SDK"],
            ["rapidocr-onnxruntime", ">=1.3.0", "OCR引擎（可选）"],
            ["wcferry", ">=39.0.0", "WeChatFerry（可选）"],
            ["pynng", ">=0.7.0", "wcferry依赖的nanomsg"],
            ["winrt-pywinrt", ">=2.0.0", "Windows系统OCR（可选）"],
            ["python-docx", ">=1.0.0", "文档生成（开发用）"],
        ],
    )

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 文档结束 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.date.today().strftime("%Y年%m月%d日")
    run = p.add_run(f"文档生成日期：{today}  |  由Claude Code自动生成")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # ── Save ──
    output_path = r"C:\Users\Trouvailler\Desktop\赵有才微信机器人_开发文档.docx"
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
